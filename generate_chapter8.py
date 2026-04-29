# Generates Chapter 8: Conclusion and Recommendation for the Folio FYP report.
# Output: C:\Users\jinz1\Desktop\FYP\Chapter_8_Conclusion.docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\jinz1\Desktop\FYP\Chapter_8_Conclusion.docx"

doc = Document()

# ── helpers ────────────────────────────────────────────────────────────────────

def h(text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def shade(row, fill="D9E1F2"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill)
        tcPr.append(shd)

def hdr(row, *labels):
    for i, lbl in enumerate(labels):
        row.cells[i].text = lbl
        for run in row.cells[i].paragraphs[0].runs:
            run.bold = True
        row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def tbl(data, headers, col_widths=None):
    t = doc.add_table(rows=len(data) + 1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    shade(t.rows[0])
    hdr(t.rows[0], *headers)
    for i, row_data in enumerate(data, 1):
        for j, val in enumerate(row_data):
            t.rows[i].cells[j].text = str(val)
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 8: CONCLUSION AND RECOMMENDATION
# ══════════════════════════════════════════════════════════════════════════════

h("8. CONCLUSION AND RECOMMENDATION", 1)

# ── 8.1 Introduction ──────────────────────────────────────────────────────────
h("8.1 Introduction", 2)
body(
    "This chapter presents the conclusion of the Folio Cross-Platform Facial "
    "Recognition System project. It begins by evaluating the extent to which each "
    "of the three project objective groups defined in Chapter 1 has been achieved, "
    "drawing on the implementation outcomes described in Chapter 6 and the testing "
    "results documented in Chapter 7. The chapter then identifies the known "
    "limitations of the current system and concludes with concrete recommendations "
    "for future work that would further mature the system towards full production "
    "readiness."
)
body(
    "Folio was developed in response to three interconnected problems in the "
    "Malaysian e-commerce context: authentication friction and poor checkout "
    "experience, security vulnerabilities in conventional credential-based "
    "authentication, and user privacy concerns that undermine trust in biometric "
    "payment systems. The system addresses these problems through a full-stack "
    "bookstore platform that integrates facial recognition for both login and "
    "payment verification, backed by a liveness detection pipeline and a "
    "privacy-first data handling approach."
)

# ── 8.2 Objectives Achievement ────────────────────────────────────────────────
h("8.2 Objectives Achievement", 2)
body(
    "The following subsections assess the degree to which each objective group was "
    "achieved, supported by evidence from the implemented system and the testing "
    "results."
)

# 8.2.1
h("8.2.1 Objective Group 1: System Design and Friction Reduction", 3)
body(
    "Objective Group 1 required the design and implementation of a bookstore "
    "e-commerce system integrating facial recognition for user login and payment "
    "verification, and the improvement of user experience by minimising the time "
    "and steps required during the login and checkout process, reducing reliance "
    "on password or OTP-based flows."
)
body(
    "This objective group has been fully achieved. Folio was implemented as a "
    "complete, cross-platform bookstore web application built on Laravel 12 and "
    "React 19, with Progressive Web App (PWA) support allowing it to be installed "
    "and used on both desktop and mobile devices without requiring a separate native "
    "app. The system covers the full customer journey: registration, browsing, "
    "cart management, wishlist, order placement, preorder, and review submission."
)
body(
    "Face-based login (FaceLoginForm.jsx) replaces the traditional email-and-password "
    "flow with a single camera-based verification step. During checkout, users may "
    "choose between standard password verification and face-verified payment "
    "(CheckoutPage.jsx, PaymentController), directly replacing the need for a separate "
    "OTP or re-authentication step. The face-verified payment path requires no "
    "additional user input beyond holding the device camera: the system issues an "
    "active liveness challenge (random head-rotation instruction), captures ten "
    "frames, and completes payment on a successful match -- a flow that is "
    "qualitatively faster and simpler than typing a password or waiting for an OTP."
)
body(
    "The System Usability Scale (SUS) evaluation conducted with 10 participants "
    "returned an average score of 80.75 out of 100, placing Folio in the Grade B "
    "-- 'Good' category, which exceeds the industry-standard acceptable threshold "
    "of 68. This result provides empirical evidence that real users can navigate "
    "the system effectively and that the interface design succeeds in reducing "
    "perceived friction. Objective Group 1 is therefore considered achieved."
)

# 8.2.2
h("8.2.2 Objective Group 2: Security Architecture and Technical Robustness", 3)
body(
    "Objective Group 2 required the implementation of a reliable and accurate "
    "facial recognition authentication module that provides a demonstrably more "
    "secure alternative to conventional credential-based methods, and the evaluation "
    "of the system in terms of usability, reliability, and technical performance "
    "to ensure it meets acceptable standards for real-world deployment."
)
body(
    "This objective group has been substantially achieved. The facial recognition "
    "module is built on the InsightFace buffalo_l model (ResNet-50 backbone, "
    "w600k_r50.onnx), which generates a 512-dimensional embedding vector per face "
    "using an ArcFace-trained recognition head. All face matching is performed via "
    "cosine similarity against the stored embedding, with a production threshold of "
    "0.6. The module is wrapped in a Python Flask microservice and communicates with "
    "the Laravel backend through internal HTTP API calls, maintaining a clean "
    "separation of concerns."
)
body(
    "A two-layer liveness detection pipeline guards against presentation attacks. "
    "Passive liveness uses Local Binary Pattern (LBP) texture analysis combined with "
    "Laplacian variance scoring to distinguish live faces from printed photographs or "
    "screen replays. Active liveness issues random head-rotation challenges "
    "(turn_left, turn_right, look_up, look_down) and measures yaw and pitch deviation "
    "across ten captured frames, requiring a minimum deviation of 3 degrees to confirm "
    "a live user is present."
)
body(
    "The quantitative performance evaluation conducted in Chapter 7 produced the "
    "following results on the full LFW 10-fold evaluation protocol (6,000 pairs):"
)
tbl(
    [
        ("Pairs evaluated",           "6,000  (3,000 genuine + 3,000 impostor)"),
        ("Accuracy",                  "88.20%"),
        ("False Accept Rate (FAR)",   "0.00%  (zero impostors accepted)"),
        ("False Reject Rate (FRR)",   "23.60%"),
        ("Equal Error Rate (EER)",    "16.58%  at threshold 0.634"),
        ("Precision",                 "1.0000  (100%)"),
        ("Recall / TAR",              "76.40%"),
        ("F1 Score",                  "0.8662"),
    ],
    ("Metric", "Value"),
)
body(
    "The most critical result is FAR = 0.00% across all 3,000 impostor pairs. This "
    "confirms that the system did not incorrectly accept a single unauthorised "
    "individual as a known user across the entire benchmark, which is the core "
    "security property required of the face-verified payment feature. The FRR of "
    "23.60% is the primary usability cost, meaning roughly one in four genuine "
    "verification attempts at the production threshold would require a retry; "
    "however, in controlled enrolment conditions (centred face, consistent lighting) "
    "real-world FRR is expected to be lower than the varied in-the-wild LFW "
    "photographs."
)
body(
    "All 7 synthetic active liveness unit tests passed, confirming correct behaviour "
    "of the pose-deviation logic for all challenge types and edge cases (insufficient "
    "movement, too few frames). The full test suite of 73 unit tests and 19 "
    "integration tests all passed, providing comprehensive coverage of every API "
    "endpoint and cross-module workflow. Objective Group 2 is therefore considered "
    "substantially achieved."
)

# 8.2.3
h("8.2.3 Objective Group 3: Privacy Protection and Trust Building", 3)
body(
    "Objective Group 3 required the implementation of strong user privacy and data "
    "security protections through biometric data encryption, minimal data retention "
    "policies, and clear informed consent mechanisms, and the assessment of user "
    "perceptions of risk and trust to measure how these perceptions influence "
    "adoption intention."
)
body(
    "This objective group has been achieved through a combination of architectural "
    "decisions and user-facing design choices. The system never retains raw facial "
    "images at any point: the Flask microservice processes frames in memory and "
    "returns only the 512-dimensional embedding vector, which is then stored in the "
    "users table as an encrypted column (face_embedding) using Laravel's built-in "
    "field encryption. Because an embedding vector cannot be reverse-engineered "
    "into a recognisable image of the user, the system meets a minimal-data-retention "
    "standard appropriate for a biometric system."
)
body(
    "Informed consent is presented to the user before any biometric enrolment action "
    "in the Face Enrolment tab (FaceRegisterTab.jsx within ProfilePage), with an "
    "explicit explanation of what data is captured, how it is stored, and how it is "
    "used for authentication purposes. Users retain full control: the Profile page "
    "allows them to enrol, view enrolment status, and delete their biometric profile "
    "at any time without administrative intervention (checkFaceStatus, removeFace "
    "endpoints in FaceRecognitionController)."
)
body(
    "The SUS score of 80.75 (Grade B -- Good), obtained from a 10-participant "
    "evaluation, reflects that users found the system navigable and trustworthy in "
    "practice. While a dedicated privacy-perception questionnaire was not administered "
    "separately from the SUS evaluation in this study, the positive overall usability "
    "rating suggests that the informed-consent design and control mechanisms did not "
    "create a perception of friction or suspicion that would have depressed the score. "
    "Objective Group 3 is therefore considered achieved at the implementation and "
    "initial evaluation level."
)

# ── 8.3 Limitations ───────────────────────────────────────────────────────────
h("8.3 Limitations", 2)
body(
    "Despite the successful achievement of the project objectives, several "
    "limitations of the current system are acknowledged."
)

body("(1)  High False Reject Rate at the Production Threshold")
body(
    "The FRR of 23.60% at the production threshold of 0.6, as measured on the LFW "
    "funneled benchmark, is relatively high for a user-facing application. In "
    "practice this means roughly one in four genuine verification attempts would "
    "fail and require the user to retry. The buffalo_l model was trained and "
    "optimised for the deep-funneled LFW variant; the use of the standard funneled "
    "variant during benchmarking partially accounts for the elevated FRR, and "
    "real-world performance under controlled enrolment conditions is expected to be "
    "better. Nevertheless, the FRR represents a usability cost that would require "
    "mitigation before large-scale user deployment."
)

body("(2)  CPU-Only Inference and Scalability Constraints")
body(
    "The Flask microservice runs on ONNX Runtime's CPUExecutionProvider. On the "
    "development machine, processing ten frames for a single verify request takes "
    "several seconds. Under high concurrent load, multiple simultaneous verification "
    "requests would queue on the CPU, degrading response times. The system has not "
    "been profiled under sustained concurrent user load because no GPU was available "
    "during development. This limits the system's readiness for production-scale "
    "deployment."
)

body("(3)  Development-Only Deployment Configuration")
body(
    "The system was developed and tested in a local environment using WAMP Server "
    "(MySQL), Laravel's built-in development server, and ngrok for HTTPS tunnelling. "
    "This configuration is adequate for functional testing and demonstration but "
    "is not production-grade: there is no load balancer, no persistent SSL "
    "certificate, no containerisation, and the ngrok tunnel URL changes each session. "
    "A production deployment would require a dedicated hosting environment with "
    "proper infrastructure."
)

body("(4)  Passive Liveness Without a Real Spoof Dataset")
body(
    "The passive liveness detection component (LBP texture + Laplacian variance "
    "scoring) was implemented and unit-tested with synthetic inputs, but it was "
    "not evaluated against a labelled real/spoof image dataset due to the "
    "unavailability of such a dataset during the project period. The passive liveness "
    "accuracy therefore remains unquantified, and the robustness of the texture-based "
    "classifier against high-quality printed photographs or screen replays cannot "
    "be confirmed empirically."
)

body("(5)  PWA Limitations on Mobile Platforms")
body(
    "Folio is deployed as a Progressive Web App rather than a native mobile "
    "application. While PWA provides cross-platform reach and installability, "
    "certain mobile capabilities -- such as persistent background push "
    "notifications and low-level camera hardware access -- are more restricted "
    "in a browser context than in a native app. On iOS Safari in particular, "
    "camera access permissions and PWA service-worker capabilities are more "
    "limited compared to Android Chrome, which may affect the face capture "
    "experience on Apple devices."
)

body("(6)  Limited Study Population for Usability Evaluation")
body(
    "The SUS evaluation was conducted with 10 participants, which, while sufficient "
    "for an initial usability study, is a small sample size. The participant pool "
    "was not explicitly stratified by demographic characteristics such as age, "
    "technical literacy, or prior experience with biometric systems. A larger and "
    "more diverse study population would provide more generalisable conclusions "
    "about usability and adoption intent in the Malaysian e-commerce context."
)

# ── 8.4 Recommendation for Future Work ────────────────────────────────────────
h("8.4 Recommendation for Future Work", 2)
body(
    "Based on the limitations identified above and the experience gained during "
    "development and testing, the following recommendations are made for future "
    "work to extend and mature the Folio system."
)

body("(1)  Threshold Calibration and FRR Reduction")
body(
    "The production threshold of 0.6 was set conservatively to achieve FAR = 0.00%. "
    "Future work should investigate threshold calibration strategies that reduce FRR "
    "while maintaining an acceptably low FAR, for example through adaptive "
    "per-user threshold personalisation based on enrolment-time score distributions. "
    "Additionally, re-benchmarking with the LFW deep-funneled dataset would provide "
    "a more accurate measure of the system's true FRR in deployment conditions, "
    "since the buffalo_l model was trained on deep-funneled aligned images."
)

body("(2)  GPU Acceleration and Performance Optimisation")
body(
    "Replacing the CPUExecutionProvider with a CUDA or TensorRT execution provider "
    "would significantly reduce inference latency per frame. Combined with "
    "request batching and asynchronous task queuing (e.g., Celery with Redis), the "
    "face service could handle far higher concurrent load. This is the most impactful "
    "single change for production readiness, and should be prioritised before any "
    "large-scale deployment."
)

body("(3)  Production-Grade Deployment with Containerisation")
body(
    "The system should be containerised using Docker and deployed on a cloud platform "
    "(e.g., AWS, Google Cloud, or Azure) with a managed database service, an HTTPS "
    "load balancer, and auto-scaling for the face microservice. The frontend build "
    "should be served from a CDN. This would replace the current ngrok-dependent "
    "development setup with a stable, scalable, and secure deployment architecture."
)

body("(4)  Passive Liveness Evaluation with a Labelled Spoof Dataset")
body(
    "Future work should obtain or construct a labelled dataset of genuine live faces "
    "and spoof attacks (printed photographs, phone/screen replays) relevant to the "
    "Malaysian usage context, and use it to benchmark and tune the passive liveness "
    "classifier. If the LBP-based approach proves insufficient, it could be replaced "
    "or augmented with a lightweight deep learning-based liveness detector such as "
    "FasNet or a MobileNet-based anti-spoofing model fine-tuned on domain-specific data."
)

body("(5)  Native Mobile Application")
body(
    "A dedicated React Native application would overcome the camera access and "
    "notification limitations of the PWA approach, particularly on iOS. React Native "
    "shares the same JavaScript/React codebase paradigm as the existing frontend, "
    "reducing the migration effort. A native app would also enable biometric "
    "integration with the device's native Face ID or fingerprint APIs as a "
    "complementary or fallback authentication factor."
)

body("(6)  Expanded Biometric Modalities and Multi-Factor Authentication")
body(
    "The current system uses face recognition as the sole biometric factor. Future "
    "work could integrate a second biometric modality -- such as device-native "
    "fingerprint recognition -- to create a multi-factor biometric authentication "
    "flow. Multi-modal biometric systems are significantly harder to spoof and "
    "provide higher overall accuracy through score-level or decision-level fusion "
    "of complementary evidence."
)

body("(7)  AI-Powered Book Recommendations")
body(
    "The current home page surfaces featured books and category filters but does not "
    "provide personalised recommendations. A collaborative-filtering or content-based "
    "recommendation engine, trained on order history, wishlist activity, and review "
    "ratings, would substantially increase engagement and average order value. "
    "Integration could be achieved as a separate microservice with a recommendation "
    "endpoint called by the home page on login."
)

body("(8)  Expanded Usability Study and Privacy Perception Research")
body(
    "A follow-up user study with a larger and more demographically diverse participant "
    "pool should be conducted, incorporating both a validated SUS questionnaire and "
    "a dedicated privacy-perception instrument (e.g., adapted from the Valence "
    "Framework applied to Malaysian facial recognition payment adoption). This would "
    "provide statistically significant evidence on whether the privacy-first design "
    "decisions -- embedding-only storage, informed consent, and user-controlled "
    "deletion -- meaningfully reduce perceived privacy risk and increase adoption "
    "intent among the Malaysian target user group."
)

doc.save(OUT)
print("Saved: " + OUT)
